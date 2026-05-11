"""Test fixtures and helpers for revisedoc tests."""

import zipfile

import pytest
from docx import Document
from lxml import etree

from revisedoc.editor import replace_text, _ns as editor_ns


# ─── Helpers ──────────────────────────────────────────────────────────────────


def read_xml(path, part="word/document.xml"):
    with zipfile.ZipFile(str(path)) as z:
        return etree.fromstring(z.read(part))


def part_names(path):
    with zipfile.ZipFile(str(path)) as z:
        return z.namelist()


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _ns(tag):
    return f"{{{W}}}{tag}"


# ─── Fixtures ──────────────────────────────────────────────────────────────────


@pytest.fixture
def simple_doc(tmp_path):
    doc = Document()
    doc.add_paragraph("Hello world foo bar")
    doc.add_paragraph("Another paragraph with foo and foo again")
    path = tmp_path / "input.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def multiline_doc(tmp_path):
    doc = Document()
    doc.add_paragraph("The quick brown fox jumps over the lazy dog near the river bank.")
    path = tmp_path / "input.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def formatted_doc(tmp_path):
    doc = Document()
    p = doc.add_paragraph()
    run_a = p.add_run("Normal text ")
    run_b = p.add_run("bold text")
    run_b.bold = True
    run_c = p.add_run(" normal again")
    path = tmp_path / "formatted.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def multi_run_doc(tmp_path):
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Hello ")
    p.add_run("World")
    p.add_run(" Foo")
    path = tmp_path / "multi_run.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def empty_doc(tmp_path):
    doc = Document()
    path = tmp_path / "empty.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def revised_doc(tmp_path):
    """A .docx that already contains a tracked change (foo -> bar)."""
    doc = Document()
    doc.add_paragraph("Hello foo world")
    replace_text(doc, "foo", "bar", author="Tester")
    path = tmp_path / "revised.docx"
    doc.save(str(path))
    return path
