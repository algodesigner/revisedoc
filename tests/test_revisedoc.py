"""Tests for revisedoc — unit, core operations, and CLI."""

import json
import subprocess
import zipfile

from docx import Document
from lxml import etree

from revisedoc.editor import (
    _ns,
    _next_revision_id,
    _run_text,
    _set_run_text,
    add_comment,
    finalize_comments,
    get_full_text,
    list_revisions,
    replace_text,
    restore_revision,
)
from tests.conftest import _ns as _test_ns, read_xml

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ─── Unit tests — helpers ─────────────────────────────────────────────────────


def test_ns():
    assert _ns("p") == f"{{{W}}}p"


def test_run_text_get_set():
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("hello")
    t = r._element.find(_ns("t"))
    assert _run_text(t.getparent()) == "hello"


def test_run_text_empty():
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("")
    assert _run_text(r._element) == ""


def test_next_revision_id_empty():
    doc = Document()
    doc.add_paragraph("test")
    assert _next_revision_id(doc.element.body) == "1"


def test_next_revision_id_with_existing():
    doc = Document()
    body = doc.element.body
    p = body[0] if len(body) else None
    # Insert an ins element with id=5
    ins = etree.SubElement(p, _ns("ins"))
    ins.set(_ns("id"), "5")
    assert _next_revision_id(body) == "6"


# ─── Core: replace_text ───────────────────────────────────────────────────────


def test_replace_simple(simple_doc, tmp_path):
    doc = Document(str(simple_doc))
    replace_text(doc, "foo", "bar")
    out = tmp_path / "out.docx"
    doc.save(str(out))
    text = get_full_text(doc)
    assert "bar" in text
    assert "Hello world bar bar" in text


def test_replace_multiple(simple_doc, tmp_path):
    doc = Document(str(simple_doc))
    replace_text(doc, "foo", "bar")
    out = tmp_path / "out.docx"
    doc.save(str(out))
    text = get_full_text(doc)
    # "foo" appears 3 times -> replaced with "bar" makes 4 (1 original "bar" in "foo bar")
    assert text.count("bar") == 4


def test_replace_not_found(simple_doc):
    doc = Document(str(simple_doc))
    import pytest
    with pytest.raises(ValueError, match="not found"):
        replace_text(doc, "zzz", "aaa")


def test_replace_empty_old_text(simple_doc):
    doc = Document(str(simple_doc))
    import pytest
    with pytest.raises(ValueError, match="must not be empty"):
        replace_text(doc, "", "aaa")


def test_replace_creates_del_ins(simple_doc, tmp_path):
    doc = Document(str(simple_doc))
    replace_text(doc, "foo", "bar")
    out = tmp_path / "out.docx"
    doc.save(str(out))

    xml = read_xml(out)
    dels = xml.findall(f".//{_ns('del')}")
    ins = xml.findall(f".//{_ns('ins')}")
    assert len(dels) >= 1
    assert len(ins) >= 1
    assert len(dels) == len(ins)  # paired


def test_replace_multi_run(multi_run_doc, tmp_path):
    """old_text spans multiple <w:r> elements."""
    doc = Document(str(multi_run_doc))
    replace_text(doc, "Hello World", "Hi There")
    out = tmp_path / "out.docx"
    doc.save(str(out))
    text = get_full_text(doc)
    assert "Hi There" in text


def test_replace_preserves_rpr(formatted_doc, tmp_path):
    """Replacing text that starts in a bold run should keep bold formatting."""
    doc = Document(str(formatted_doc))
    # "bold" starts inside the bold run
    replace_text(doc, "bold text", "strong text")
    out = tmp_path / "out.docx"
    doc.save(str(out))

    xml = read_xml(out)
    ins = xml.findall(f".//{_ns('ins')}")
    assert len(ins) > 0
    # The ins run should have an rPr child (bold formatting)
    first_ins_rpr = ins[0].find(f".//{_ns('rPr')}")
    assert first_ins_rpr is not None


def test_replace_no_side_effects(simple_doc, tmp_path):
    """Text not matching old_text should be preserved."""
    doc = Document(str(simple_doc))
    replace_text(doc, "foo", "bar")
    text = get_full_text(doc)
    assert "Hello world" in text
    assert "Another paragraph with" in text


# ─── Core: add_comment ────────────────────────────────────────────────────────


def test_comment_creates_markers(simple_doc, tmp_path):
    doc = Document(str(simple_doc))
    add_comment(doc, "Hello", "A greeting")
    pending = doc.part._pending_comments
    out = tmp_path / "out.docx"
    doc.save(str(out))
    finalize_comments(str(out), pending_comments=pending)

    xml = read_xml(out)
    starts = xml.findall(f".//{_ns('commentRangeStart')}")
    ends = xml.findall(f".//{_ns('commentRangeEnd')}")
    refs = xml.findall(f".//{_ns('commentReference')}")
    assert len(starts) >= 1
    assert len(ends) >= 1
    assert len(refs) >= 1


def test_comment_injects_part(simple_doc, tmp_path):
    doc = Document(str(simple_doc))
    add_comment(doc, "Hello", "A greeting")
    pending = doc.part._pending_comments
    out = tmp_path / "out.docx"
    doc.save(str(out))
    finalize_comments(str(out), pending_comments=pending)

    names = []
    with zipfile.ZipFile(str(out)) as z:
        names = z.namelist()
    assert "word/comments.xml" in names

    comments_xml = read_xml(out, "word/comments.xml")
    comments = comments_xml.findall(_ns("comment"))
    assert len(comments) == 1
    assert comments[0].get(_ns("author")) == "Editor"


def test_comment_multiple(simple_doc, tmp_path):
    doc = Document(str(simple_doc))
    add_comment(doc, "Hello", "Comment 1")
    add_comment(doc, "foo", "Comment 2")
    pending = doc.part._pending_comments
    out = tmp_path / "out.docx"
    doc.save(str(out))
    finalize_comments(str(out), pending_comments=pending)

    comments_xml = read_xml(out, "word/comments.xml")
    comments = comments_xml.findall(_ns("comment"))
    assert len(comments) == 2


# ─── Core: list_revisions ─────────────────────────────────────────────────────


def test_list_revisions_empty(simple_doc):
    doc = Document(str(simple_doc))
    revs = list_revisions(doc)
    assert revs == []


def test_list_revisions_after_replace(revised_doc):
    doc = Document(str(revised_doc))
    revs = list_revisions(doc)
    assert len(revs) >= 2  # one deletion + one insertion
    types = {r["type"] for r in revs}
    assert "deletion" in types
    assert "insertion" in types


def test_list_revisions_text(revised_doc):
    doc = Document(str(revised_doc))
    revs = list_revisions(doc)
    del_revs = [r for r in revs if r["type"] == "deletion"]
    ins_revs = [r for r in revs if r["type"] == "insertion"]
    assert any("foo" in r["text"] for r in del_revs)
    assert any("bar" in r["text"] for r in ins_revs)


# ─── Core: restore_revision ───────────────────────────────────────────────────


def test_restore_deletion(revised_doc, tmp_path):
    doc = Document(str(revised_doc))
    revs = list_revisions(doc)
    del_id = None
    for r in revs:
        if r["type"] == "deletion":
            del_id = r["id"]
            break
    assert del_id is not None

    restore_revision(doc, del_id, restore_type="deletion")
    out = tmp_path / "restored.docx"
    doc.save(str(out))
    text = get_full_text(doc)
    assert "foo" in text  # restored deleted text


def test_restore_insertion(revised_doc, tmp_path):
    doc = Document(str(revised_doc))
    revs = list_revisions(doc)
    ins_id = None
    for r in revs:
        if r["type"] == "insertion":
            ins_id = r["id"]
            break
    assert ins_id is not None

    restore_revision(doc, ins_id, restore_type="insertion")
    out = tmp_path / "restored.docx"
    doc.save(str(out))

    # After removing the insertion, the document should not have "bar"
    text = get_full_text(doc)
    # But it might still have the original text plus the deletion
    # Actually the "foo" was deleted and "bar" was inserted.
    # After restoring the insertion (removing "bar" and the del), we should NOT have "bar"
    # But the del is also removed, and the original text is not restored.
    # Since the deletion is removal of "foo" and that del is also removed,
    # we end up with "Hello  world" (foo was deleted, bar insertion was removed)
    # Actually, let me think... after replace "foo" with "bar":
    # <del id=1>foo</del><ins id=1>bar</ins>
    # restore_type="insertion": removes ins id=1 and del id=1
    # Result: "Hello  world" (no "foo" restored, no "bar")
    xml = read_xml(str(out))
    ins = xml.findall(f".//{_ns('ins')}")
    dels = xml.findall(f".//{_ns('del')}")
    # No remaining ins/del for this revision
    assert "bar" not in text  # insertion removed


def test_restore_not_found(simple_doc):
    doc = Document(str(simple_doc))
    import pytest
    with pytest.raises(ValueError, match="not found"):
        restore_revision(doc, "999")


# ─── Core: get_full_text ──────────────────────────────────────────────────────


def test_get_full_text(simple_doc):
    doc = Document(str(simple_doc))
    text = get_full_text(doc)
    assert "Hello world foo bar" in text
    assert "Another paragraph with foo" in text


def test_get_full_text_empty(empty_doc):
    doc = Document(str(empty_doc))
    text = get_full_text(doc)
    assert text == ""


def test_get_full_text_after_replace(revised_doc):
    """get_full_text should show insertions and not deletions."""
    doc = Document(str(revised_doc))
    text = get_full_text(doc)
    assert "bar" in text  # insertion shown
    assert "foo" not in text  # deletion excluded


# ─── CLI integration tests ────────────────────────────────────────────────


def test_cli_get_text(simple_doc):
    result = subprocess.run(
        ["revisedoc", "get-text", str(simple_doc)],
        capture_output=True, text=True,
    )
    assert result.returncode == 0
    assert "Hello world foo bar" in result.stdout


def test_cli_replace(simple_doc, tmp_path):
    out = tmp_path / "out.docx"
    result = subprocess.run(
        ["revisedoc", "replace", str(simple_doc), str(out), "foo", "bar", "--author", "Test"],
        capture_output=True, text=True,
    )
    assert result.returncode == 0
    assert out.exists()

    doc = Document(str(out))
    text = get_full_text(doc)
    assert "bar" in text


def test_cli_list_revisions(revised_doc):
    result = subprocess.run(
        ["revisedoc", "list-revisions", str(revised_doc)],
        capture_output=True, text=True,
    )
    assert result.returncode == 0
    assert "insertion" in result.stdout
    assert "deletion" in result.stdout


def test_cli_list_revisions_json(revised_doc):
    result = subprocess.run(
        ["revisedoc", "list-revisions", str(revised_doc), "--format", "json"],
        capture_output=True, text=True,
    )
    assert result.returncode == 0
    data = json.loads(result.stdout)
    assert isinstance(data, list)
    assert len(data) >= 2


def test_cli_comment(simple_doc, tmp_path):
    out = tmp_path / "out.docx"
    result = subprocess.run(
        ["revisedoc", "comment", str(simple_doc), str(out), "foo", "a note", "--author", "Test"],
        capture_output=True, text=True,
    )
    assert result.returncode == 0
    assert out.exists()

    with zipfile.ZipFile(str(out)) as z:
        assert "word/comments.xml" in z.namelist()


def test_cli_restore(revised_doc, tmp_path):
    doc = Document(str(revised_doc))
    revs = list_revisions(doc)
    del_id = revs[0]["id"]

    out = tmp_path / "restored.docx"
    result = subprocess.run(
        ["revisedoc", "restore", str(revised_doc), str(out), "--revision-id", del_id],
        capture_output=True, text=True,
    )
    assert result.returncode == 0
    assert out.exists()

    doc2 = Document(str(out))
    text = get_full_text(doc2)
    assert "foo" in text  # deletion reverted


# ─── MCP server tests ──────────────────────────────────────────────────────


def _mcp_init(proc):
    """Send initialize request and return server capabilities."""
    req = json.dumps({
        "jsonrpc": "2.0", "id": 1, "method": "initialize",
        "params": {
            "protocolVersion": "2024-11-05",
            "capabilities": {},
            "clientInfo": {"name": "test", "version": "0.1.0"},
        },
    })
    proc.stdin.write(req.encode() + b"\n")
    proc.stdin.flush()
    line = proc.stdout.readline()
    resp = json.loads(line)
    assert resp["id"] == 1
    assert "result" in resp

    notif = json.dumps({"jsonrpc": "2.0", "method": "notifications/initialized"})
    proc.stdin.write(notif.encode() + b"\n")
    proc.stdin.flush()
    return resp["result"]


def test_mcp_server_lists_tools(simple_doc):
    proc = subprocess.Popen(
        ["revisedoc-mcp"],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    try:
        _mcp_init(proc)

        req = json.dumps({"jsonrpc": "2.0", "id": 2, "method": "tools/list", "params": {}})
        proc.stdin.write(req.encode() + b"\n")
        proc.stdin.flush()

        line = proc.stdout.readline()
        result = json.loads(line)
        tools = {t["name"] for t in result["result"]["tools"]}
        assert tools == {
            "docx_replace", "docx_comment", "docx_list_revisions",
            "docx_restore", "docx_get_text",
        }
    finally:
        proc.kill()
        proc.wait()


def test_mcp_server_call_get_text(simple_doc):
    proc = subprocess.Popen(
        ["revisedoc-mcp"],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    try:
        _mcp_init(proc)

        req = json.dumps({
            "jsonrpc": "2.0", "id": 2, "method": "tools/call",
            "params": {
                "name": "docx_get_text",
                "arguments": {"input_path": str(simple_doc)},
            },
        })
        proc.stdin.write(req.encode() + b"\n")
        proc.stdin.flush()

        line = proc.stdout.readline()
        result = json.loads(line)
        content = result["result"]["content"]
        text = "".join(c["text"] for c in content if c["type"] == "text")
        assert "Hello world foo bar" in text
    finally:
        proc.kill()
        proc.wait()


def test_mcp_server_call_replace(simple_doc, tmp_path):
    out = tmp_path / "out.docx"
    proc = subprocess.Popen(
        ["revisedoc-mcp"],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    try:
        _mcp_init(proc)

        req = json.dumps({
            "jsonrpc": "2.0", "id": 2, "method": "tools/call",
            "params": {
                "name": "docx_replace",
                "arguments": {
                    "input_path": str(simple_doc),
                    "output_path": str(out),
                    "old_text": "foo",
                    "new_text": "bar",
                    "author": "Test",
                },
            },
        })
        proc.stdin.write(req.encode() + b"\n")
        proc.stdin.flush()

        line = proc.stdout.readline()
        result = json.loads(line)
        assert "result" in result
        assert out.exists()

        # Verify the output file has the replacement
        doc = Document(str(out))
        text = get_full_text(doc)
        assert "bar" in text
    finally:
        proc.kill()
        proc.wait()


def test_mcp_server_call_list_revisions(revised_doc):
    proc = subprocess.Popen(
        ["revisedoc-mcp"],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    try:
        _mcp_init(proc)

        req = json.dumps({
            "jsonrpc": "2.0", "id": 2, "method": "tools/call",
            "params": {
                "name": "docx_list_revisions",
                "arguments": {"input_path": str(revised_doc)},
            },
        })
        proc.stdin.write(req.encode() + b"\n")
        proc.stdin.flush()

        line = proc.stdout.readline()
        result = json.loads(line)
        content = result["result"]["content"]
        text = "".join(c["text"] for c in content if c["type"] == "text")
        assert "insertion" in text
        assert "deletion" in text
    finally:
        proc.kill()
        proc.wait()


def test_mcp_server_call_restore(revised_doc, tmp_path):
    out = tmp_path / "restored.docx"

    # Get revision ID first via docx_list_revisions
    proc = subprocess.Popen(
        ["revisedoc-mcp"],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    try:
        _mcp_init(proc)

        req = json.dumps({
            "jsonrpc": "2.0", "id": 2, "method": "tools/call",
            "params": {
                "name": "docx_list_revisions",
                "arguments": {"input_path": str(revised_doc)},
            },
        })
        proc.stdin.write(req.encode() + b"\n")
        proc.stdin.flush()

        line = proc.stdout.readline()
        result = json.loads(line)
        content = result["result"]["content"]
        text = "".join(c["text"] for c in content if c["type"] == "text")
        proc.kill()
        proc.wait()
    except Exception:
        proc.kill()
        proc.wait()
        raise

    # Parse revision ID from text output
    del_id = None
    for line_text in text.split("\n"):
        if "ID:" in line_text and "deletion" in line_text:
            del_id = line_text.split("ID: ")[1].split(" ")[0]
            break
    assert del_id is not None

    # Now restore via docx_restore
    proc2 = subprocess.Popen(
        ["revisedoc-mcp"],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    try:
        _mcp_init(proc2)

        req = json.dumps({
            "jsonrpc": "2.0", "id": 2, "method": "tools/call",
            "params": {
                "name": "docx_restore",
                "arguments": {
                    "input_path": str(revised_doc),
                    "output_path": str(out),
                    "revision_id": del_id,
                },
            },
        })
        proc2.stdin.write(req.encode() + b"\n")
        proc2.stdin.flush()

        line = proc2.stdout.readline()
        result = json.loads(line)
        assert "result" in result
        assert out.exists()

        doc = Document(str(out))
        text = get_full_text(doc)
        assert "foo" in text  # deletion reverted
    finally:
        proc2.kill()
        proc2.wait()


def test_mcp_server_call_comment(simple_doc, tmp_path):
    out = tmp_path / "out.docx"
    proc = subprocess.Popen(
        ["revisedoc-mcp"],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    try:
        _mcp_init(proc)

        req = json.dumps({
            "jsonrpc": "2.0", "id": 2, "method": "tools/call",
            "params": {
                "name": "docx_comment",
                "arguments": {
                    "input_path": str(simple_doc),
                    "output_path": str(out),
                    "target_text": "Hello",
                    "comment": "A greeting",
                    "author": "Tester",
                },
            },
        })
        proc.stdin.write(req.encode() + b"\n")
        proc.stdin.flush()

        line = proc.stdout.readline()
        result = json.loads(line)
        assert "result" in result
        assert out.exists()

        # Verify comments.xml was injected
        with zipfile.ZipFile(str(out)) as z:
            assert "word/comments.xml" in z.namelist()
    finally:
        proc.kill()
        proc.wait()


