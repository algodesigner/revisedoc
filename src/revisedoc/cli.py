"""
Command-line interface for revisedoc.
"""

import argparse
import json

from docx import Document

from revisedoc.editor import (
    add_comment,
    finalize_comments,
    get_full_text,
    list_revisions,
    replace_text,
    restore_revision,
)


def main():
    parser = argparse.ArgumentParser(
        description="Edit .docx files with change tracking, comments, and revision restoration."
    )
    parser.add_argument("--author", default="Editor", help="Author name for tracked changes")
    sub = parser.add_subparsers(dest="command", required=True)

    replace_p = sub.add_parser("replace", help="Replace text with tracked changes")
    replace_p.add_argument("input", help="Input .docx file")
    replace_p.add_argument("output", help="Output .docx file")
    replace_p.add_argument("old_text", help="Text to replace")
    replace_p.add_argument("new_text", help="Replacement text")
    replace_p.add_argument("--author", default="Editor")

    comment_p = sub.add_parser("comment", help="Add a comment to text")
    comment_p.add_argument("input", help="Input .docx file")
    comment_p.add_argument("output", help="Output .docx file")
    comment_p.add_argument("target", help="Text to annotate")
    comment_p.add_argument("comment", help="Comment text")
    comment_p.add_argument("--author", default="Editor")

    list_p = sub.add_parser("list-revisions", help="List tracked changes")
    list_p.add_argument("input", help="Input .docx file")
    list_p.add_argument("--format", choices=["text", "json"], default="text", help="Output format")

    restore_p = sub.add_parser("restore", help="Restore a revision by ID")
    restore_p.add_argument("input", help="Input .docx file")
    restore_p.add_argument("output", help="Output .docx file")
    restore_p.add_argument("--revision-id", required=True, help="Revision ID to restore")
    restore_p.add_argument("--restore-type", choices=["deletion", "insertion"],
                           default="deletion", help="What to restore (default: deletion)")

    text_p = sub.add_parser("get-text", help="Print document plain text")
    text_p.add_argument("input", help="Input .docx file")

    args = parser.parse_args()

    if args.command == "list-revisions":
        doc = Document(args.input)
        revs = list_revisions(doc)
        if args.format == "json":
            print(json.dumps(revs, indent=2))
        else:
            if not revs:
                print("No revisions found.")
                return
            for r in revs:
                print(f"ID: {r['id']} | Type: {r['type']} | Author: {r['author']}")
                print(f"  Text: {r['text'][:80]}")
                print(f"  In paragraph: {r['paragraph_text'][:80]}")
                print()
        return

    elif args.command == "get-text":
        doc = Document(args.input)
        print(get_full_text(doc))
        return

    elif args.command == "replace":
        doc = Document(args.input)
        replace_text(doc, args.old_text, args.new_text, author=args.author)
        doc.save(args.output)
        print(f"Saved to {args.output}")
        return

    elif args.command == "comment":
        doc = Document(args.input)
        add_comment(doc, args.target, args.comment, author=args.author)
        pending = getattr(doc.part, "_pending_comments", None)
        doc.save(args.output)
        finalize_comments(args.output, pending_comments=pending)
        print(f"Saved to {args.output}")
        return

    elif args.command == "restore":
        doc = Document(args.input)
        restore_revision(doc, args.revision_id, restore_type=args.restore_type)
        doc.save(args.output)
        print(f"Saved to {args.output}")
        return


if __name__ == "__main__":
    main()
