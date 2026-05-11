"""
MCP server providing revisedoc operations as structured tools.

Start with:  revisedoc-mcp
"""

from docx import Document
from mcp.server.fastmcp import FastMCP

from revisedoc.editor import (
    add_comment,
    finalize_comments,
    get_full_text,
    inspect_document,
    list_revisions,
    replace_text,
    restore_revision,
    search_text,
)

server = FastMCP("revisedoc")


@server.tool(description="""Replace text in a .docx file with Word-compatible tracked changes. The old text is wrapped in <w:del> (strikethrough) and the new text in <w:ins> (underline), exactly as if a human used Word's "Track Changes" feature. All occurrences are replaced. Use this when you need to update a contract, proposal, or any document while preserving a visible edit history.""")
def docx_replace(
    input_path: str,
    output_path: str,
    old_text: str,
    new_text: str,
    author: str = "Editor",
) -> str:
    doc = Document(input_path)
    replace_text(doc, old_text, new_text, author=author)
    doc.save(output_path)
    return f"Replaced {old_text!r} with {new_text!r} -> {output_path}"


@server.tool(description="""Add an inline comment anchored to a specific text range in a .docx file. The comment appears in Word's comment pane and the annotated text gets highlighted. Use this for document review workflows — flagging sections for revision, asking questions, or leaving notes for collaborators.""")
def docx_comment(
    input_path: str,
    output_path: str,
    target_text: str,
    comment: str,
    author: str = "Editor",
) -> str:
    doc = Document(input_path)
    add_comment(doc, target_text, comment, author=author)
    pending = getattr(doc.part, "_pending_comments", None)
    doc.save(output_path)
    finalize_comments(output_path, pending_comments=pending)
    return f"Added comment on {target_text!r} -> {output_path}"


@server.tool(description="""List all tracked changes (insertions and deletions) in a .docx file. Returns each revision's ID, type, author, date, and text content. Use this BEFORE restoring a revision to discover the revision IDs, or to inspect what changes a document contains. The "json" format is best for programmatic use.""")
def docx_list_revisions(
    input_path: str,
    format: str = "text",
) -> str:
    import json
    doc = Document(input_path)
    revs = list_revisions(doc)
    if format == "json":
        return json.dumps(revs, indent=2)
    if not revs:
        return "No revisions found."
    lines = []
    for r in revs:
        lines.append(f"ID: {r['id']} | Type: {r['type']} | Author: {r['author']}")
        lines.append(f"  Text: {r['text'][:80]}")
        lines.append(f"  In paragraph: {r['paragraph_text'][:80]}")
        lines.append("")
    return "\n".join(lines)


@server.tool(description="""Undo a specific tracked change in a .docx file by its revision ID. Use docx_list_revisions first to discover IDs. With restore_type="deletion" (default), the deleted text is moved back and the paired insertion is removed — restoring the original text. With restore_type="insertion", both the insertion and its paired deletion are removed, effectively discarding the edit.""")
def docx_restore(
    input_path: str,
    output_path: str,
    revision_id: str,
    restore_type: str = "deletion",
) -> str:
    doc = Document(input_path)
    restore_revision(doc, revision_id, restore_type=restore_type)
    doc.save(output_path)
    return f"Restored revision {revision_id} -> {output_path}"


@server.tool(description="""Extract the plain text of a .docx file with insertions applied and deletions excluded. Returns paragraphs joined by newlines with no paragraph indices or formatting info. For structured output with paragraph indices and per-run formatting, use docx_inspect instead.""")
def docx_get_text(
    input_path: str,
) -> str:
    doc = Document(input_path)
    return get_full_text(doc)


@server.tool(description="""Inspect the structure of a .docx file. Returns every paragraph with its index, paraId, effective text, run-level formatting (bold/italic/underline), and revision counts. Optionally pass a paragraph_index to inspect a single paragraph in detail. Use this BEFORE editing to understand the document structure and formatting.""")
def docx_inspect(
    input_path: str,
    paragraph_index: int | None = None,
) -> str:
    import json
    doc = Document(input_path)
    info = inspect_document(doc, paragraph_index=paragraph_index)
    return json.dumps(info, indent=2)


@server.tool(description="""Search for exact text in a .docx file. Returns all occurrences with the paragraph index, character offsets (match_start, match_end) within the paragraph, and which runs the match spans. Use this BEFORE docx_replace to verify the text exists and understand its exact location in the document structure. Returns an empty list if no matches are found.""")
def docx_search(
    input_path: str,
    query: str,
) -> str:
    import json
    doc = Document(input_path)
    results = search_text(doc, query)
    return json.dumps(results, indent=2)


def main():
    server.run(transport="stdio")


if __name__ == "__main__":
    main()
